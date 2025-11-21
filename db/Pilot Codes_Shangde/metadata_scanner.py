"""
Minimal Metadata Scanner for Urban Planning Data
Scans directories for data files and extracts standardized metadata without loading full datasets.
"""

import os
import json
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
import mimetypes


class DataAnalyzer:
    """Base class for data analysis"""

    @staticmethod
    def get_file_size(file_path: str) -> Dict[str, Any]:
        """Get file size in bytes and human-readable format"""
        size_bytes = os.path.getsize(file_path)
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return {'bytes': os.path.getsize(file_path), 'readable': f"{size_bytes:.2f} {unit}"}
            size_bytes /= 1024.0
        return {'bytes': os.path.getsize(file_path), 'readable': f"{size_bytes:.2f} PB"}


class GeospatialAnalyzer(DataAnalyzer):
    """Analyzer for geospatial data formats"""

    SUPPORTED_FORMATS = ['.shp', '.geojson', '.json', '.tif', '.tiff', '.kml', '.kmz', '.gdb']

    @staticmethod
    def analyze(file_path: str) -> Dict[str, Any]:
        """Analyze geospatial data without loading full dataset"""
        metadata = {
            'data_type': 'geospatial',
            'file_size': DataAnalyzer.get_file_size(file_path),
            'format': Path(file_path).suffix.lower(),
            'features': [],
            'sample_data': {},
            'applicable_operations': []
        }

        ext = Path(file_path).suffix.lower()

        try:
            # Vector formats
            if ext in ['.shp', '.geojson', '.json', '.kml']:
                metadata.update(GeospatialAnalyzer._analyze_vector(file_path))
            # Raster formats
            elif ext in ['.tif', '.tiff']:
                metadata.update(GeospatialAnalyzer._analyze_raster(file_path))
            # Geodatabase
            elif ext == '.gdb':
                metadata.update(GeospatialAnalyzer._analyze_geodatabase(file_path))

        except Exception as e:
            metadata['error'] = f"Analysis failed: {str(e)}"
            metadata['applicable_operations'] = ['Manual inspection required']

        return metadata

    @staticmethod
    def _convert_bounds_to_latlon(bounds: tuple, source_crs) -> Dict[str, Any]:
        """
        Convert bounding box from source CRS to WGS84 (lat/lon)

        Args:
            bounds: Tuple of (minx, miny, maxx, maxy) in source CRS
            source_crs: Source coordinate reference system

        Returns:
            Dictionary with lat/lon bounds and formatted string
        """
        try:
            from pyproj import Transformer, CRS

            # Check if already in lat/lon (EPSG:4326 or similar)
            source_crs_obj = CRS.from_user_input(source_crs)
            if source_crs_obj.is_geographic:
                # Already in lat/lon
                return {
                    'west': bounds[0],
                    'south': bounds[1],
                    'east': bounds[2],
                    'north': bounds[3],
                    'formatted': f"West: {bounds[0]:.6f}, South: {bounds[1]:.6f}, East: {bounds[2]:.6f}, North: {bounds[3]:.6f}"
                }

            # Transform to WGS84 (EPSG:4326)
            transformer = Transformer.from_crs(source_crs, "EPSG:4326", always_xy=True)

            # Transform corners
            min_lon, min_lat = transformer.transform(bounds[0], bounds[1])
            max_lon, max_lat = transformer.transform(bounds[2], bounds[3])

            return {
                'west': min_lon,
                'south': min_lat,
                'east': max_lon,
                'north': max_lat,
                'formatted': f"West: {min_lon:.6f}, South: {min_lat:.6f}, East: {max_lon:.6f}, North: {max_lat:.6f}"
            }

        except ImportError:
            return {
                'error': 'pyproj not installed. Install with: pip install pyproj',
                'formatted': 'Conversion requires pyproj library'
            }
        except Exception as e:
            return {
                'error': f'Conversion failed: {str(e)}',
                'formatted': 'Unable to convert bounds'
            }

    @staticmethod
    def _analyze_vector(file_path: str) -> Dict[str, Any]:
        """Analyze vector data (Shapefile, GeoJSON, KML) using sampling"""
        result = {}

        try:
            import fiona

            with fiona.open(file_path, 'r') as src:
                # Get schema without loading features
                result['crs'] = src.crs.to_string() if src.crs else 'Unknown'
                result['geometry_type'] = src.schema['geometry']
                result['feature_count'] = len(src)
                result['bounds'] = src.bounds

                # Convert bounds to lat/lon for easier interpretation
                if src.crs:
                    result['bounds_latlon'] = GeospatialAnalyzer._convert_bounds_to_latlon(
                        src.bounds, src.crs
                    )

                # Get field names and types
                properties = src.schema['properties']
                result['features'] = [
                    {'name': name, 'type': dtype}
                    for name, dtype in properties.items()
                ]

                # Determine applicable geoprocessing operations based on geometry type
                # Do this BEFORE sampling so operations are always generated even if sampling fails
                operations = GeospatialAnalyzer._get_vector_operations(
                    result['geometry_type'],
                    properties
                )
                result['applicable_operations'] = operations

                # Sample first few features only
                # Wrapped in try-except so sampling errors don't break the entire analysis
                try:
                    sample_size = min(5, len(src))
                    samples = {}
                    for idx, feature in enumerate(src):
                        if idx >= sample_size:
                            break
                        for prop_name, prop_value in feature['properties'].items():
                            if prop_name not in samples:
                                samples[prop_name] = []
                            # Skip None/null values to avoid errors
                            if prop_value is not None:
                                samples[prop_name].append(prop_value)

                    result['sample_data'] = {k: list(set(v)) for k, v in samples.items()}
                except Exception as sample_error:
                    # If sampling fails, just skip it and continue
                    result['sample_data'] = {}
                    result['sampling_warning'] = f"Could not sample data: {str(sample_error)}"

        except ImportError:
            result['error'] = 'Fiona not installed. Install with: pip install fiona'
            result['applicable_operations'] = ['Install fiona/geopandas for analysis']
        except Exception as e:
            result['error'] = f"Vector analysis failed: {str(e)}"

        return result

    @staticmethod
    def _get_vector_operations(geometry_type: str, properties: Dict[str, str]) -> List[str]:
        """Determine applicable operations based on geometry type and attributes"""
        operations = []

        # Common spatial operations for all vector types
        operations.extend([
            'Spatial query/filter',
            'Spatial join (with other layers)',
            'Clip/Intersect (with boundary)',
            'Buffer analysis',
            'Reproject/Transform CRS',
            'Export to different format'
        ])

        # Geometry-specific operations
        if 'Polygon' in geometry_type:
            operations.extend([
                'Calculate area',
                'Dissolve by attribute',
                'Union/Merge polygons',
                'Simplify geometry',
                'Calculate centroid',
                'Zonal statistics (with raster)'
            ])
        elif 'LineString' in geometry_type or 'Line' in geometry_type:
            operations.extend([
                'Calculate length',
                'Merge/Dissolve lines',
                'Split lines at intersections',
                'Create points along line',
                'Network analysis'
            ])
        elif 'Point' in geometry_type:
            operations.extend([
                'Point density/heatmap',
                'Nearest neighbor analysis',
                'Cluster analysis',
                'Voronoi diagram',
                'Thiessen polygons'
            ])

        # Attribute-based operations
        categorical_attrs = [name for name, dtype in properties.items() if 'str' in dtype.lower()]
        numerical_attrs = [name for name, dtype in properties.items() if 'int' in dtype.lower() or 'float' in dtype.lower()]

        if categorical_attrs:
            operations.append(f'Dissolve/Group by categorical attributes: {", ".join(categorical_attrs[:3])}')
            operations.append(f'Symbolize by categories: {", ".join(categorical_attrs[:3])}')

        if numerical_attrs:
            operations.append(f'Calculate statistics for: {", ".join(numerical_attrs[:3])}')
            operations.append(f'Choropleth mapping by: {", ".join(numerical_attrs[:3])}')

        return operations

    @staticmethod
    def _analyze_raster(file_path: str) -> Dict[str, Any]:
        """Analyze raster data (GeoTIFF) without loading full array"""
        result = {}

        try:
            import rasterio

            with rasterio.open(file_path) as src:
                result['crs'] = src.crs.to_string() if src.crs else 'Unknown'
                result['dimensions'] = {'width': src.width, 'height': src.height}
                result['band_count'] = src.count
                result['dtype'] = str(src.dtypes[0])
                result['bounds'] = src.bounds
                result['resolution'] = src.res
                result['nodata_value'] = src.nodata

                # Convert bounds to lat/lon for easier interpretation
                if src.crs:
                    result['bounds_latlon'] = GeospatialAnalyzer._convert_bounds_to_latlon(
                        src.bounds, src.crs
                    )

                # Sample small window instead of full raster
                window = rasterio.windows.Window(0, 0, min(100, src.width), min(100, src.height))
                sample = src.read(1, window=window)

                result['sample_data'] = {
                    'min': float(sample.min()),
                    'max': float(sample.max()),
                    'mean': float(sample.mean())
                }

                result['features'] = [f'Band_{i+1}' for i in range(src.count)]

                # Determine applicable raster operations
                operations = GeospatialAnalyzer._get_raster_operations(src.count, str(src.dtypes[0]))
                result['applicable_operations'] = operations

        except ImportError:
            result['error'] = 'Rasterio not installed. Install with: pip install rasterio'
            result['applicable_operations'] = ['Install rasterio for analysis']
        except Exception as e:
            result['error'] = f"Raster analysis failed: {str(e)}"

        return result

    @staticmethod
    def _get_raster_operations(band_count: int, dtype: str) -> List[str]:
        """Determine applicable raster operations"""
        operations = [
            'Resample/Change resolution',
            'Reproject to different CRS',
            'Clip/Extract by mask',
            'Raster calculator/Band math',
            'Extract values to points',
            'Reclassify values',
            'Export to different format'
        ]

        # Multi-band specific operations
        if band_count > 1:
            operations.extend([
                'Composite bands (RGB visualization)',
                'Band ratio/index calculation (e.g., NDVI)',
                'Principal Component Analysis (PCA)',
                'Change detection (multi-temporal)'
            ])

        # Continuous data operations (float/int)
        if 'float' in dtype.lower() or 'int' in dtype.lower():
            operations.extend([
                'Calculate slope/aspect',
                'Generate contours',
                'Hillshade analysis',
                'Zonal statistics',
                'Viewshed analysis',
                'Cost-distance analysis',
                'Interpolation/Kriging'
            ])

        return operations

    @staticmethod
    def _analyze_geodatabase(file_path: str) -> Dict[str, Any]:
        """Analyze file geodatabase"""
        result = {
            'error': 'File Geodatabase analysis requires ArcPy or GDAL',
            'applicable_operations': ['Requires specialized tools for analysis']
        }
        return result


class TabularAnalyzer(DataAnalyzer):
    """Analyzer for tabular data formats"""

    SUPPORTED_FORMATS = ['.csv', '.xlsx', '.xls', '.parquet', '.tsv', '.txt']

    @staticmethod
    def analyze(file_path: str) -> Dict[str, Any]:
        """Analyze tabular data using sampling"""
        metadata = {
            'data_type': 'tabular',
            'file_size': DataAnalyzer.get_file_size(file_path),
            'format': Path(file_path).suffix.lower(),
            'features': [],
            'sample_data': {},
            'applicable_operations': []
        }

        try:
            import pandas as pd

            ext = Path(file_path).suffix.lower()

            # Read only first N rows for sampling
            sample_size = 100

            if ext == '.csv':
                df = pd.read_csv(file_path, nrows=sample_size)
            elif ext in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path, nrows=sample_size)
            elif ext == '.parquet':
                df = pd.read_parquet(file_path)[:sample_size]
            elif ext in ['.tsv', '.txt']:
                df = pd.read_csv(file_path, sep='\t', nrows=sample_size)
            else:
                raise ValueError(f"Unsupported format: {ext}")

            metadata['row_count_sample'] = len(df)
            metadata['column_count'] = len(df.columns)

            # Analyze each column
            features = []
            sample_data = {}

            for col in df.columns:
                col_info = {
                    'name': col,
                    'dtype': str(df[col].dtype),
                    'null_count': int(df[col].isnull().sum()),
                    'null_percentage': round(df[col].isnull().sum() / len(df) * 100, 2)
                }

                # Categorical or numerical analysis
                if df[col].dtype in ['object', 'category']:
                    unique_values = df[col].dropna().unique()
                    col_info['unique_count'] = len(unique_values)
                    sample_data[col] = list(unique_values[:10])  # First 10 unique values
                elif df[col].dtype in ['int64', 'float64', 'int32', 'float32']:
                    col_info['min'] = float(df[col].min()) if not df[col].isnull().all() else None
                    col_info['max'] = float(df[col].max()) if not df[col].isnull().all() else None
                    col_info['mean'] = float(df[col].mean()) if not df[col].isnull().all() else None
                    sample_data[col] = [float(df[col].iloc[i]) for i in range(min(5, len(df))) if not pd.isna(df[col].iloc[i])]

                features.append(col_info)

            metadata['features'] = features
            metadata['sample_data'] = sample_data

            # Determine applicable tabular operations
            operations = TabularAnalyzer._get_tabular_operations(features, df)
            metadata['applicable_operations'] = operations

        except ImportError:
            metadata['error'] = 'Pandas not installed. Install with: pip install pandas'
            metadata['applicable_operations'] = ['Install pandas for analysis']
        except Exception as e:
            metadata['error'] = f"Tabular analysis failed: {str(e)}"
            metadata['applicable_operations'] = ['Manual inspection required']

        return metadata

    @staticmethod
    def _get_tabular_operations(features: List[Dict], df) -> List[str]:
        """Determine applicable tabular operations based on data characteristics"""
        operations = []

        # Common operations for all tabular data
        operations.extend([
            'Filter/Query rows',
            'Sort by column',
            'Remove duplicates',
            'Export to different format'
        ])

        # Identify column types
        categorical_cols = [f['name'] for f in features if f['dtype'] in ['object', 'category']]
        numerical_cols = [f['name'] for f in features if 'int' in f['dtype'] or 'float' in f['dtype']]
        cols_with_nulls = [f['name'] for f in features if f['null_count'] > 0]

        # Categorical operations
        if categorical_cols:
            operations.append(f'Group by/Aggregate on: {", ".join(categorical_cols[:3])}')
            operations.append(f'Pivot table by: {", ".join(categorical_cols[:2])}')
            operations.append(f'Value counts for: {", ".join(categorical_cols[:3])}')
            operations.append(f'One-hot encoding for: {", ".join(categorical_cols[:3])}')

        # Numerical operations
        if numerical_cols:
            operations.append(f'Calculate statistics (mean, median, std) for: {", ".join(numerical_cols[:3])}')
            operations.append(f'Binning/Categorization of: {", ".join(numerical_cols[:3])}')
            operations.append(f'Normalize/Standardize: {", ".join(numerical_cols[:3])}')
            operations.append(f'Correlation analysis between numerical columns')

        # Missing value handling
        if cols_with_nulls:
            operations.append(f'Fill missing values in: {", ".join(cols_with_nulls[:3])}')
            operations.append(f'Drop rows/columns with missing values')

        # Multi-column operations
        if len(df.columns) >= 2:
            operations.extend([
                'Join/Merge with other tables',
                'Concatenate with other tables',
                'Create calculated/derived columns'
            ])

        # Time series operations if datetime columns detected
        datetime_cols = [col for col in df.columns if 'date' in col.lower() or 'time' in col.lower()]
        if datetime_cols:
            operations.append(f'Time series analysis on: {", ".join(datetime_cols[:2])}')
            operations.append(f'Temporal aggregation (daily, monthly, yearly)')

        # Spatial join if lat/lon columns exist
        lat_cols = [col for col in df.columns if 'lat' in col.lower() or 'y' in col.lower()]
        lon_cols = [col for col in df.columns if 'lon' in col.lower() or 'x' in col.lower()]
        if lat_cols and lon_cols:
            operations.append(f'Convert to geospatial data using: {lat_cols[0]}, {lon_cols[0]}')
            operations.append('Spatial join with geographic layers')

        return operations


class DocumentAnalyzer(DataAnalyzer):
    """Analyzer for document formats"""

    SUPPORTED_FORMATS = ['.pdf', '.docx', '.doc', '.txt', '.md']

    @staticmethod
    def analyze(file_path: str) -> Dict[str, Any]:
        """Analyze document files"""
        metadata = {
            'data_type': 'document',
            'file_size': DataAnalyzer.get_file_size(file_path),
            'format': Path(file_path).suffix.lower(),
            'features': [],
            'sample_data': {},
            'applicable_operations': []
        }

        ext = Path(file_path).suffix.lower()

        try:
            if ext == '.pdf':
                metadata.update(DocumentAnalyzer._analyze_pdf(file_path))
            elif ext == '.docx':
                metadata.update(DocumentAnalyzer._analyze_docx(file_path))
            elif ext in ['.txt', '.md']:
                metadata.update(DocumentAnalyzer._analyze_text(file_path))
            else:
                metadata['error'] = f"Unsupported document format: {ext}"

        except Exception as e:
            metadata['error'] = f"Document analysis failed: {str(e)}"

        return metadata

    @staticmethod
    def _analyze_pdf(file_path: str) -> Dict[str, Any]:
        """Analyze PDF document"""
        result = {}

        try:
            import PyPDF2

            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                result['page_count'] = len(reader.pages)

                # Extract text from first page only
                first_page_text = reader.pages[0].extract_text()
                result['sample_data'] = {
                    'first_page_preview': first_page_text[:500] + '...' if len(first_page_text) > 500 else first_page_text
                }

                result['features'] = ['text_content', 'pages']
                result['applicable_operations'] = [
                    'Extract full text content',
                    'Extract text from specific pages',
                    'Keyword search and extraction',
                    'Text summarization (LLM)',
                    'Entity extraction (locations, organizations, dates)',
                    'Topic modeling',
                    'Convert to structured data (tables)',
                    'Merge with other PDFs'
                ]

        except ImportError:
            result['error'] = 'PyPDF2 not installed. Install with: pip install PyPDF2'
            result['applicable_operations'] = ['Install PyPDF2 for analysis']
        except Exception as e:
            result['error'] = f"PDF analysis failed: {str(e)}"

        return result

    @staticmethod
    def _analyze_docx(file_path: str) -> Dict[str, Any]:
        """Analyze DOCX document"""
        result = {}

        try:
            import docx

            doc = docx.Document(file_path)
            result['paragraph_count'] = len(doc.paragraphs)
            result['table_count'] = len(doc.tables)

            # Extract first few paragraphs
            preview = '\n'.join([p.text for p in doc.paragraphs[:3] if p.text])
            result['sample_data'] = {
                'preview': preview[:500] + '...' if len(preview) > 500 else preview
            }

            result['features'] = ['text_content', 'paragraphs', 'tables']
            result['applicable_operations'] = [
                'Extract full text content',
                'Extract tables to CSV/Excel',
                'Keyword search and extraction',
                'Text summarization (LLM)',
                'Entity extraction (locations, organizations, dates)',
                'Section/heading extraction',
                'Convert to PDF',
                'Merge with other documents'
            ]

        except ImportError:
            result['error'] = 'python-docx not installed. Install with: pip install python-docx'
            result['applicable_operations'] = ['Install python-docx for analysis']
        except Exception as e:
            result['error'] = f"DOCX analysis failed: {str(e)}"

        return result

    @staticmethod
    def _analyze_text(file_path: str) -> Dict[str, Any]:
        """Analyze plain text document"""
        result = {}

        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                # Read only first 10KB
                content = file.read(10240)
                lines = content.split('\n')

                result['line_count_sample'] = len(lines)
                result['sample_data'] = {
                    'preview': content[:500] + '...' if len(content) > 500 else content
                }

                result['features'] = ['text_content']
                result['applicable_operations'] = [
                    'Full text search',
                    'Keyword extraction',
                    'Text summarization (LLM)',
                    'Entity extraction (locations, organizations, dates)',
                    'Sentiment analysis',
                    'Topic modeling',
                    'Convert to structured format (JSON, CSV)'
                ]

        except Exception as e:
            result['error'] = f"Text analysis failed: {str(e)}"

        return result


class MetadataScanner:
    """Main scanner class that orchestrates directory scanning and metadata extraction"""

    def __init__(self):
        self.analyzers = {
            'geospatial': GeospatialAnalyzer,
            'tabular': TabularAnalyzer,
            'document': DocumentAnalyzer
        }

    def scan_directory(self, directory_path: str, output_file: Optional[str] = None) -> Dict[str, Any]:
        """
        Scan directory and extract metadata for all supported data files

        Args:
            directory_path: Path to directory to scan
            output_file: Optional path to save JSON output

        Returns:
            Dictionary containing metadata for all discovered files
        """
        directory_path = Path(directory_path)

        if not directory_path.exists():
            raise ValueError(f"Directory not found: {directory_path}")

        metadata = {
            'scan_timestamp': datetime.now().isoformat(),
            'scan_directory': str(directory_path.absolute()),
            'files': []
        }

        # Find all supported files
        supported_extensions = (
            GeospatialAnalyzer.SUPPORTED_FORMATS +
            TabularAnalyzer.SUPPORTED_FORMATS +
            DocumentAnalyzer.SUPPORTED_FORMATS
        )

        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                print(f"Analyzing: {file_path.name}")
                file_metadata = self._analyze_file(str(file_path))
                metadata['files'].append(file_metadata)

        metadata['total_files'] = len(metadata['files'])

        # Save to JSON if output file specified
        if output_file:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False)
            print(f"\nMetadata saved to: {output_file}")

        return metadata

    def _analyze_file(self, file_path: str) -> Dict[str, Any]:
        """Analyze a single file and return metadata"""
        file_metadata = {
            'file_path': str(Path(file_path).absolute()),
            'file_name': Path(file_path).name,
            'file_extension': Path(file_path).suffix.lower(),
            'last_modified': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
        }

        # Determine analyzer type
        ext = Path(file_path).suffix.lower()

        if ext in GeospatialAnalyzer.SUPPORTED_FORMATS:
            analysis = GeospatialAnalyzer.analyze(file_path)
        elif ext in TabularAnalyzer.SUPPORTED_FORMATS:
            analysis = TabularAnalyzer.analyze(file_path)
        elif ext in DocumentAnalyzer.SUPPORTED_FORMATS:
            analysis = DocumentAnalyzer.analyze(file_path)
        else:
            analysis = {
                'data_type': 'unknown',
                'error': 'Unsupported file type'
            }

        file_metadata.update(analysis)
        return file_metadata


def main():
    """Main entry point for the scanner"""
    import argparse

    parser = argparse.ArgumentParser(description='Scan directory and extract metadata from data files')
    parser.add_argument('directory', help='Directory path to scan')
    parser.add_argument('-o', '--output', help='Output JSON file path', default='metadata_output.json')

    args = parser.parse_args()

    scanner = MetadataScanner()

    print(f"Scanning directory: {args.directory}")
    print("=" * 60)

    metadata = scanner.scan_directory(args.directory, args.output)

    print("=" * 60)
    print(f"Scan complete! Found {metadata['total_files']} files.")
    print(f"Metadata saved to: {args.output}")


if __name__ == '__main__':
    main()
